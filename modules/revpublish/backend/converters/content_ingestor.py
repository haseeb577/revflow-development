"""
RevPublish Content Ingestor
Supports: DOCX, Google Docs, Google Sheets, JSON, HTML, CSV
Converts content to page type field mappings for deployment
"""

import os
import re
import json
import tempfile
import subprocess
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Optional dependencies
try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    DOCX_NATIVE = True
except ImportError:
    DOCX_NATIVE = False
    logger.info("python-docx not available, using pandoc fallback")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    logger.warning("BeautifulSoup not available, HTML parsing limited")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas not available, spreadsheet parsing limited")


class ContentIngestor:
    """
    Ingests content from multiple formats and maps to RevPublish page types.

    Supported formats:
    - DOCX (MS Word)
    - HTML
    - JSON (structured or raw)
    - CSV (Google Sheets export)
    - Markdown
    - Plain text
    """

    # Field mapping patterns for auto-detection
    FIELD_PATTERNS = {
        'hero_headline': ['headline', 'title', 'h1', 'main heading', 'hero'],
        'hero_subheadline': ['subheadline', 'subtitle', 'tagline', 'subheading'],
        'service_name': ['service', 'service name', 'offering'],
        'service_description': ['description', 'about', 'overview', 'details'],
        'benefits': ['benefits', 'advantages', 'why choose', 'features'],
        'phone_number': ['phone', 'tel', 'call', 'telephone'],
        'email': ['email', 'mail', 'contact email'],
        'address': ['address', 'location', 'street'],
        'city': ['city', 'town', 'municipality'],
        'state': ['state', 'province', 'region'],
        'content': ['content', 'body', 'text', 'article'],
        'excerpt': ['excerpt', 'summary', 'intro', 'teaser'],
        'cta_text': ['cta', 'call to action', 'button text'],
        'company_name': ['company', 'business name', 'organization'],
        'price': ['price', 'cost', 'fee', 'rate'],
        'product_name': ['product', 'item', 'name'],
    }

    def __init__(self):
        self.supported_formats = ['.docx', '.html', '.htm', '.json', '.csv', '.md', '.txt']

    def ingest_file(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """
        Ingest a file and return structured data for RevPublish.

        Args:
            file_path: Path to the file
            page_type_id: Optional page type to map to

        Returns:
            Dict with 'data' (field mappings) and 'metadata'
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.docx':
            return self.ingest_docx(file_path, page_type_id)
        elif suffix in ['.html', '.htm']:
            return self.ingest_html(file_path, page_type_id)
        elif suffix == '.json':
            return self.ingest_json(file_path, page_type_id)
        elif suffix == '.csv':
            return self.ingest_csv(file_path, page_type_id)
        elif suffix == '.md':
            return self.ingest_markdown(file_path, page_type_id)
        elif suffix == '.txt':
            return self.ingest_text(file_path, page_type_id)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def ingest_content(self, content: str, content_type: str, page_type_id: Optional[str] = None) -> Dict:
        """
        Ingest raw content string.

        Args:
            content: The content string
            content_type: 'html', 'json', 'markdown', 'text'
            page_type_id: Optional page type to map to
        """
        if content_type == 'html':
            return self._parse_html(content, page_type_id)
        elif content_type == 'json':
            return self._parse_json(content, page_type_id)
        elif content_type in ['markdown', 'md']:
            return self._parse_markdown(content, page_type_id)
        else:
            return self._parse_text(content, page_type_id)

    # === DOCX Ingestion ===

    def ingest_docx(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest MS Word DOCX file"""

        if DOCX_NATIVE:
            return self._ingest_docx_native(file_path, page_type_id)
        else:
            return self._ingest_docx_pandoc(file_path, page_type_id)

    def _ingest_docx_native(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse DOCX using python-docx library"""
        doc = DocxDocument(file_path)

        extracted = {
            'paragraphs': [],
            'headings': [],
            'lists': [],
            'tables': []
        }

        current_heading = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ''

            if 'Heading' in style or style.startswith('Title'):
                level = 1
                if 'Heading 2' in style:
                    level = 2
                elif 'Heading 3' in style:
                    level = 3
                extracted['headings'].append({'text': text, 'level': level})
                current_heading = text
            elif 'List' in style or para.text.startswith(('•', '-', '*', '1.', '2.')):
                extracted['lists'].append({
                    'text': text,
                    'under_heading': current_heading
                })
            else:
                extracted['paragraphs'].append({
                    'text': text,
                    'under_heading': current_heading
                })

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                extracted['tables'].append(table_data)

        return self._map_to_page_type(extracted, page_type_id, 'docx')

    def _ingest_docx_pandoc(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse DOCX using pandoc (fallback)"""
        try:
            result = subprocess.run(
                ['pandoc', file_path, '-t', 'html'],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0:
                return self._parse_html(result.stdout, page_type_id)
            else:
                raise Exception(f"Pandoc error: {result.stderr}")
        except FileNotFoundError:
            raise Exception("Pandoc not installed. Install with: apt-get install pandoc")

    # === HTML Ingestion ===

    def ingest_html(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest HTML file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._parse_html(content, page_type_id)

    def _parse_html(self, html_content: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse HTML content to extract structured data"""

        if BS4_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')

            extracted = {
                'paragraphs': [],
                'headings': [],
                'lists': [],
                'tables': [],
                'raw_html': html_content
            }

            # Extract title
            title_tag = soup.find('title')
            if title_tag:
                extracted['title'] = title_tag.get_text().strip()

            # Extract headings
            for level in range(1, 7):
                for h in soup.find_all(f'h{level}'):
                    extracted['headings'].append({
                        'text': h.get_text().strip(),
                        'level': level
                    })

            # Extract paragraphs
            for p in soup.find_all('p'):
                text = p.get_text().strip()
                if text:
                    extracted['paragraphs'].append({'text': text})

            # Extract lists
            for ul in soup.find_all(['ul', 'ol']):
                items = [li.get_text().strip() for li in ul.find_all('li')]
                extracted['lists'].append({'items': items})

            # Extract tables
            for table in soup.find_all('table'):
                table_data = []
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if cells:
                        table_data.append(cells)
                if table_data:
                    extracted['tables'].append(table_data)

            return self._map_to_page_type(extracted, page_type_id, 'html')
        else:
            # Basic regex parsing without BeautifulSoup
            extracted = {
                'raw_html': html_content,
                'paragraphs': [],
                'headings': []
            }

            # Extract title
            title_match = re.search(r'<title>(.*?)</title>', html_content, re.IGNORECASE | re.DOTALL)
            if title_match:
                extracted['title'] = title_match.group(1).strip()

            # Extract h1
            h1_match = re.search(r'<h1[^>]*>(.*?)</h1>', html_content, re.IGNORECASE | re.DOTALL)
            if h1_match:
                extracted['headings'].append({'text': re.sub(r'<[^>]+>', '', h1_match.group(1)).strip(), 'level': 1})

            return self._map_to_page_type(extracted, page_type_id, 'html')

    # === JSON Ingestion ===

    def ingest_json(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._parse_json(content, page_type_id)

    def _parse_json(self, json_content: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse JSON content"""
        try:
            data = json.loads(json_content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON: {e}")

        # If it's already structured with our field names, use directly
        if isinstance(data, dict):
            # Check if it matches our page type fields
            if page_type_id and any(key in self.FIELD_PATTERNS for key in data.keys()):
                return {
                    'data': data,
                    'metadata': {
                        'source_format': 'json',
                        'page_type_id': page_type_id,
                        'direct_mapping': True
                    }
                }

            # Try to map fields
            extracted = {'raw_data': data}
            for key, value in data.items():
                if isinstance(value, str):
                    extracted[key] = value
                elif isinstance(value, list) and all(isinstance(i, str) for i in value):
                    extracted[key] = '|'.join(value)  # Convert to pipe-separated

            return self._map_to_page_type(extracted, page_type_id, 'json')

        elif isinstance(data, list):
            # Array of items - could be multiple pages
            return {
                'data': {'items': data},
                'metadata': {
                    'source_format': 'json',
                    'is_array': True,
                    'item_count': len(data)
                },
                'multiple_items': data
            }

        return {'data': data, 'metadata': {'source_format': 'json'}}

    # === CSV/Google Sheets Ingestion ===

    def ingest_csv(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest CSV file (exported from Google Sheets)"""

        if PANDAS_AVAILABLE:
            return self._ingest_csv_pandas(file_path, page_type_id)
        else:
            return self._ingest_csv_native(file_path, page_type_id)

    def _ingest_csv_pandas(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse CSV using pandas"""
        df = pd.read_csv(file_path)

        # Clean column names
        df.columns = [col.strip().lower().replace(' ', '_') for col in df.columns]

        records = df.to_dict('records')

        if len(records) == 1:
            # Single row - treat as single page
            return self._map_to_page_type({'raw_data': records[0], **records[0]}, page_type_id, 'csv')
        else:
            # Multiple rows - batch import
            return {
                'data': {'items': records},
                'metadata': {
                    'source_format': 'csv',
                    'row_count': len(records),
                    'columns': list(df.columns),
                    'is_batch': True
                },
                'multiple_items': [
                    self._map_to_page_type({'raw_data': row, **row}, page_type_id, 'csv')
                    for row in records
                ]
            }

    def _ingest_csv_native(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse CSV without pandas"""
        import csv

        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            records = list(reader)

        # Clean keys
        cleaned_records = []
        for row in records:
            cleaned = {k.strip().lower().replace(' ', '_'): v for k, v in row.items()}
            cleaned_records.append(cleaned)

        if len(cleaned_records) == 1:
            return self._map_to_page_type({'raw_data': cleaned_records[0], **cleaned_records[0]}, page_type_id, 'csv')
        else:
            return {
                'data': {'items': cleaned_records},
                'metadata': {
                    'source_format': 'csv',
                    'row_count': len(cleaned_records),
                    'is_batch': True
                },
                'multiple_items': [
                    self._map_to_page_type({'raw_data': row, **row}, page_type_id, 'csv')
                    for row in cleaned_records
                ]
            }

    # === Markdown Ingestion ===

    def ingest_markdown(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._parse_markdown(content, page_type_id)

    def _parse_markdown(self, md_content: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse Markdown content"""
        extracted = {
            'paragraphs': [],
            'headings': [],
            'lists': [],
            'raw_markdown': md_content
        }

        lines = md_content.split('\n')
        current_heading = None
        current_list = []

        for line in lines:
            stripped = line.strip()

            if stripped.startswith('# '):
                extracted['headings'].append({'text': stripped[2:], 'level': 1})
                current_heading = stripped[2:]
            elif stripped.startswith('## '):
                extracted['headings'].append({'text': stripped[3:], 'level': 2})
                current_heading = stripped[3:]
            elif stripped.startswith('### '):
                extracted['headings'].append({'text': stripped[4:], 'level': 3})
                current_heading = stripped[4:]
            elif stripped.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', stripped):
                # List item
                list_text = re.sub(r'^[-*+]\s*|\d+\.\s*', '', stripped)
                current_list.append(list_text)
            elif stripped and current_list:
                # End of list
                extracted['lists'].append({
                    'items': current_list,
                    'under_heading': current_heading
                })
                current_list = []
                extracted['paragraphs'].append({
                    'text': stripped,
                    'under_heading': current_heading
                })
            elif stripped:
                extracted['paragraphs'].append({
                    'text': stripped,
                    'under_heading': current_heading
                })

        # Don't forget last list
        if current_list:
            extracted['lists'].append({
                'items': current_list,
                'under_heading': current_heading
            })

        return self._map_to_page_type(extracted, page_type_id, 'markdown')

    # === Plain Text Ingestion ===

    def ingest_text(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """Ingest plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._parse_text(content, page_type_id)

    def _parse_text(self, text_content: str, page_type_id: Optional[str] = None) -> Dict:
        """Parse plain text content"""
        lines = text_content.strip().split('\n')

        extracted = {
            'paragraphs': [],
            'raw_text': text_content
        }

        # First non-empty line could be title
        for line in lines:
            if line.strip():
                extracted['title'] = line.strip()
                break

        # Rest is content
        content_lines = [l for l in lines[1:] if l.strip()]
        extracted['content'] = '\n'.join(content_lines)

        return self._map_to_page_type(extracted, page_type_id, 'text')

    # === Field Mapping ===

    def _map_to_page_type(self, extracted: Dict, page_type_id: Optional[str], source_format: str) -> Dict:
        """
        Map extracted content to page type fields.
        Uses pattern matching to auto-detect field mappings.
        """
        mapped_data = {}

        # Direct field mappings from extracted data
        if 'raw_data' in extracted:
            for key, value in extracted['raw_data'].items():
                # Try to map key to known field
                mapped_field = self._match_field(key)
                if mapped_field:
                    mapped_data[mapped_field] = value
                else:
                    mapped_data[key] = value

        # Extract from headings
        if 'headings' in extracted and extracted['headings']:
            h1_headings = [h for h in extracted['headings'] if h.get('level') == 1]
            if h1_headings:
                mapped_data['hero_headline'] = h1_headings[0]['text']

            h2_headings = [h for h in extracted['headings'] if h.get('level') == 2]
            if h2_headings and 'hero_subheadline' not in mapped_data:
                mapped_data['hero_subheadline'] = h2_headings[0]['text']

        # Extract from title
        if 'title' in extracted and 'hero_headline' not in mapped_data:
            mapped_data['hero_headline'] = extracted['title']

        # Extract content from paragraphs
        if 'paragraphs' in extracted:
            content_parts = [p['text'] for p in extracted['paragraphs']]
            if content_parts:
                # First paragraph could be description/intro
                if 'service_description' not in mapped_data and 'content' not in mapped_data:
                    mapped_data['service_description'] = '<p>' + '</p><p>'.join(content_parts) + '</p>'
                if 'excerpt' not in mapped_data and content_parts:
                    mapped_data['excerpt'] = content_parts[0][:200]

        # Extract from lists (benefits, features, etc.)
        if 'lists' in extracted:
            for list_item in extracted['lists']:
                items = list_item.get('items', [])
                heading = list_item.get('under_heading', '').lower()

                if items:
                    pipe_list = '|'.join(items)

                    # Try to map based on heading context
                    if 'benefit' in heading or 'why' in heading:
                        mapped_data['benefits'] = pipe_list
                    elif 'feature' in heading:
                        mapped_data['features'] = pipe_list
                    elif 'step' in heading or 'process' in heading:
                        mapped_data['process_steps'] = pipe_list
                    elif 'area' in heading or 'service' in heading:
                        mapped_data['service_areas'] = pipe_list
                    elif 'benefits' not in mapped_data:
                        # Default to benefits if no specific match
                        mapped_data['benefits'] = pipe_list

        # Keep raw HTML if available
        if 'raw_html' in extracted:
            mapped_data['_raw_html'] = extracted['raw_html']
        if 'raw_markdown' in extracted:
            mapped_data['_raw_markdown'] = extracted['raw_markdown']

        return {
            'data': mapped_data,
            'metadata': {
                'source_format': source_format,
                'page_type_id': page_type_id,
                'extracted_fields': list(mapped_data.keys()),
                'auto_mapped': True
            }
        }

    def _match_field(self, input_key: str) -> Optional[str]:
        """Try to match an input key to a known field name"""
        input_lower = input_key.lower().replace('_', ' ').replace('-', ' ')

        for field_name, patterns in self.FIELD_PATTERNS.items():
            for pattern in patterns:
                if pattern in input_lower or input_lower in pattern:
                    return field_name

        return None


class GoogleDocsIngestor:
    """
    Handles Google Docs and Google Sheets via exported files or API.

    For API access, requires credentials in:
    /opt/revpublish/config/google_credentials.json
    """

    def __init__(self):
        self.credentials_path = '/opt/revpublish/config/google_credentials.json'
        self.base_ingestor = ContentIngestor()
        self._api_available = False
        self._init_api()

    def _init_api(self):
        """Initialize Google API if credentials available"""
        try:
            from google.oauth2.service_account import Credentials
            from googleapiclient.discovery import build

            if os.path.exists(self.credentials_path):
                self.credentials = Credentials.from_service_account_file(
                    self.credentials_path,
                    scopes=[
                        'https://www.googleapis.com/auth/documents.readonly',
                        'https://www.googleapis.com/auth/spreadsheets.readonly',
                        'https://www.googleapis.com/auth/drive.readonly'
                    ]
                )
                self._api_available = True
                logger.info("Google API initialized")
        except ImportError:
            logger.info("Google API libraries not installed. Use exported files.")
        except Exception as e:
            logger.warning(f"Could not initialize Google API: {e}")

    def ingest_google_doc(self, doc_id: str, page_type_id: Optional[str] = None) -> Dict:
        """
        Ingest a Google Doc by document ID.

        Args:
            doc_id: Google Docs document ID (from URL)
            page_type_id: Target page type
        """
        if not self._api_available:
            raise Exception("Google API not configured. Export as DOCX and use file import.")

        from googleapiclient.discovery import build

        service = build('docs', 'v1', credentials=self.credentials)
        doc = service.documents().get(documentId=doc_id).execute()

        # Extract content from Google Docs structure
        extracted = {
            'title': doc.get('title', ''),
            'headings': [],
            'paragraphs': [],
            'lists': []
        }

        content = doc.get('body', {}).get('content', [])

        for element in content:
            if 'paragraph' in element:
                para = element['paragraph']
                text = ''
                for elem in para.get('elements', []):
                    if 'textRun' in elem:
                        text += elem['textRun'].get('content', '')

                text = text.strip()
                if not text:
                    continue

                # Check paragraph style
                style = para.get('paragraphStyle', {}).get('namedStyleType', '')

                if 'HEADING' in style:
                    level = int(style.replace('HEADING_', '')) if style != 'HEADING' else 1
                    extracted['headings'].append({'text': text, 'level': level})
                elif 'TITLE' in style:
                    extracted['headings'].append({'text': text, 'level': 1})
                else:
                    extracted['paragraphs'].append({'text': text})

        return self.base_ingestor._map_to_page_type(extracted, page_type_id, 'google_doc')

    def ingest_google_sheet(self, spreadsheet_id: str, sheet_name: Optional[str] = None,
                           page_type_id: Optional[str] = None) -> Dict:
        """
        Ingest a Google Sheet by spreadsheet ID.

        Args:
            spreadsheet_id: Google Sheets spreadsheet ID
            sheet_name: Optional specific sheet name (defaults to first sheet)
            page_type_id: Target page type
        """
        if not self._api_available:
            raise Exception("Google API not configured. Export as CSV and use file import.")

        from googleapiclient.discovery import build

        service = build('sheets', 'v4', credentials=self.credentials)

        # Get spreadsheet metadata
        spreadsheet = service.spreadsheets().get(spreadsheetId=spreadsheet_id).execute()

        # Determine range
        if sheet_name:
            range_name = f"'{sheet_name}'"
        else:
            range_name = spreadsheet['sheets'][0]['properties']['title']

        # Get values
        result = service.spreadsheets().values().get(
            spreadsheetId=spreadsheet_id,
            range=range_name
        ).execute()

        values = result.get('values', [])

        if not values:
            return {'data': {}, 'metadata': {'source_format': 'google_sheet', 'empty': True}}

        # First row is headers
        headers = [h.strip().lower().replace(' ', '_') for h in values[0]]

        records = []
        for row in values[1:]:
            record = {}
            for i, value in enumerate(row):
                if i < len(headers):
                    record[headers[i]] = value
            records.append(record)

        if len(records) == 1:
            return self.base_ingestor._map_to_page_type(
                {'raw_data': records[0], **records[0]},
                page_type_id,
                'google_sheet'
            )
        else:
            return {
                'data': {'items': records},
                'metadata': {
                    'source_format': 'google_sheet',
                    'row_count': len(records),
                    'columns': headers,
                    'is_batch': True
                },
                'multiple_items': [
                    self.base_ingestor._map_to_page_type(
                        {'raw_data': row, **row},
                        page_type_id,
                        'google_sheet'
                    )
                    for row in records
                ]
            }

    def ingest_exported_file(self, file_path: str, page_type_id: Optional[str] = None) -> Dict:
        """
        Ingest an exported Google Docs/Sheets file.
        Accepts: .docx (Google Docs export), .csv/.xlsx (Sheets export)
        """
        return self.base_ingestor.ingest_file(file_path, page_type_id)


# Singleton instances
content_ingestor = ContentIngestor()
google_ingestor = GoogleDocsIngestor()
